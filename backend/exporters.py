from datetime import datetime
from io import BytesIO

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def export_to_pdf(content: str, title: str = "Document Summary") -> bytes:
    if not REPORTLAB_AVAILABLE:
        raise Exception("PDF export not available")
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        textColor='#333333'
    )
    
    story = []
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 12))
    
    lines = content.split('\n')
    for line in lines:
        if line.strip():
            if line.startswith('#'):
                story.append(Paragraph(line.replace('#', '').strip(), styles['Heading2']))
            elif line.startswith('•') or line.startswith('-'):
                story.append(Paragraph(line, styles['Bullet']))
            else:
                story.append(Paragraph(line, styles['Normal']))
            story.append(Spacer(1, 6))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def export_to_docx(content: str, title: str = "Document Summary") -> bytes:
    if not DOCX_AVAILABLE:
        raise Exception("DOCX export not available")
    
    doc = Document()
    doc.add_heading(title, 0)
    
    lines = content.split('\n')
    for line in lines:
        if line.strip():
            if line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            elif line.startswith('##'):
                doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('•') or line.startswith('-'):
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(line.replace('•', '').replace('-', '').strip())
            else:
                doc.add_paragraph(line)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def export_to_markdown(content: str, title: str = "Document Summary") -> str:
    markdown_content = f"# {title}\n\n"
    markdown_content += f"*Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}*\n\n"
    
    lines = content.split('\n')
    for line in lines:
        if line.strip():
            if line.startswith('•'):
                markdown_content += f"- {line[1:].strip()}\n"
            elif 'Key Points:' in line or 'SUMMARY' in line.upper():
                markdown_content += f"## {line.strip()}\n\n"
            elif line.endswith(':') and len(line.split()) <= 3:
                markdown_content += f"### {line.strip()}\n\n"
            else:
                markdown_content += f"{line.strip()}\n\n"
    
    return markdown_content

def export_to_txt(content: str, title: str = "Document Summary") -> str:
    txt_content = f"{title}\n{'=' * len(title)}\n\n"
    txt_content += f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n\n"
    txt_content += content
    
    return txt_content